from django.shortcuts import render
from blog.models import Entry, Category, Bookmark, BookmarkCategory, AppCategory, \
    AppEntry
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from tagging.models import Tag
from django.core.mail import send_mail

from bs4 import BeautifulSoup

from docx import Document

import xlwt
from django.http import StreamingHttpResponse

import requests


# Create your views here.
def index(request):
    latest_entry_list = Entry.objects.order_by('creation_date')[:10]
    context = {'latest_entry_list': latest_entry_list}
    return render(request, 'index.html', context)


def blog(request):
    latest_entry_list = Entry.objects.order_by('-creation_date')
    paginator = Paginator(latest_entry_list, 10)
    page = request.GET.get('page')

    try:
        entrys = paginator.page(page)
    except PageNotAnInteger:
        entrys = paginator.page(1)
    except EmptyPage:
        entrys = paginator.page(paginator.num_pages)

    catagory_list = Category.objects.raw(
        "select cate.id, cate.title, cate.description, count(cate.id) as entrynum from blog_category as cate, blog_entry_categorys as enca where cate.id = enca.category_id group by cate.id")

    tag_list = Tag.objects.all()

    context = {'latest_entry_list': entrys, 'catagory_list': catagory_list,
               'latest_post': latest_entry_list[:3], 'tag_list': tag_list}

    return render(request, 'blog.html', context)


def single(request, post_id):
    # post_id = request.GET.get('id')
    post = Entry.objects.get(id=post_id)

    latest_entry_list = Entry.objects.order_by('-creation_date')[:3]
    catagory_list = Category.objects.raw(
        "select cate.id, cate.title, cate.description, count(cate.id) as entrynum from blog_category as cate, blog_entry_categorys as enca where cate.id = enca.category_id group by cate.id")
    tag_list = Tag.objects.all()

    context = {'post': post, 'catagory_list': catagory_list,
               'latest_post': latest_entry_list, 'tag_list': tag_list}
    return render(request, 'single.html', context)


def category(request, category_id):
    cate_entry_list = Entry.objects.filter(
        categorys__id=category_id).order_by('-creation_date')

    paginator = Paginator(cate_entry_list, 10)
    page = request.GET.get('page')

    try:
        entrys = paginator.page(page)
    except PageNotAnInteger:
        entrys = paginator.page(1)
    except EmptyPage:
        entrys = paginator.page(paginator.num_pages)

    latest_entry_list = Entry.objects.order_by('-creation_date')[:3]
    catagory_list = Category.objects.raw(
        "select cate.id, cate.title, cate.description, count(cate.id) as entrynum from blog_category as cate, blog_entry_categorys as enca where cate.id = enca.category_id group by cate.id")
    tag_list = Tag.objects.all()

    context = {'latest_entry_list': entrys, 'catagory_list': catagory_list,
               'latest_post': latest_entry_list, 'tag_list': tag_list}

    return render(request, 'blog.html', context)


def taglist(request, tag_id):

    tag = Tag.objects.get(id=tag_id)
    tag_entry_list = Entry.objects.filter(tags__icontains=tag.name)
    paginator = Paginator(tag_entry_list, 10)
    page = request.GET.get('page')

    try:
        entrys = paginator.page(page)
    except PageNotAnInteger:
        entrys = paginator.page(1)
    except EmptyPage:
        entrys = paginator.page(paginator.num_pages)

    latest_entry_list = Entry.objects.order_by('-creation_date')[:3]
    catagory_list = Category.objects.raw(
        "select cate.id, cate.title, cate.description, count(cate.id) as entrynum from blog_category as cate, blog_entry_categorys as enca where cate.id = enca.category_id group by cate.id")
    tag_list = Tag.objects.all()

    context = {'latest_entry_list': entrys, 'catagory_list': catagory_list,
               'latest_post': latest_entry_list, 'tag_list': tag_list}

    return render(request, 'blog.html', context)


def bookmarklist(request):

    bookmarks = Bookmark.objects.raw(
        "select book.id, book.title, book.description, book.click_times, book.url, book.head_image, cate.title as category_title from blog_bookmark as book, blog_bookmarkcategory as cate where book.categorys_id = cate.id and book.status = 2")

    categorys = BookmarkCategory.objects.all()
    context = {'bookmark_list': bookmarks, 'category_list': categorys}
    return render(request, 'bookmark.html', context)


def appcategory(request, appcategory_id):
    appcategory = AppCategory.objects.get(id=appcategory_id)
    applist = AppEntry.objects.filter(categorys_id=appcategory_id)
    context = {'appcategory': appcategory, 'applist': applist}
    return render(request, 'service.html', context)


def contact(request):
    latest_entry_list = Entry.objects.order_by('-creation_date')[:3]
    catagory_list = Category.objects.raw(
        "select cate.id, cate.title, cate.description, count(cate.id) as entrynum from blog_category as cate, blog_entry_categorys as enca where cate.id = enca.category_id group by cate.id")
    tag_list = Tag.objects.all()
    context = {'catagory_list': catagory_list,
               'latest_post': latest_entry_list, 'tag_list': tag_list}

    return render(request, 'contact.html', context)


def sendmail(request):
    name = request.POST.get('name', 'customer')
    email = request.POST.get('email', 'melonblogs@163.com')
    subject = request.POST.get('subject')
    msg = request.POST.get('msg')
    msg = 'name:' + name + 'from email:' + email + 'subject:' + msg

    send_mail(subject, msg, 'melonblogs@163.com', [
              '59170121@qq.com'], fail_silently=False)
    return render(request, 'contact.html')


def crawler(request):
    print "Hello Beautiful!"

    base_url = "http://www.searchcio.com.cn/"
    campaigns = "news.htm"

    print "Starting at url: " + base_url + campaigns

    # Fetch page from the web and get the text response
    r = requests.get(base_url + campaigns)
    data = r.text

    print "Retrived data, getting campaign links."
    # Parse it into a format we can handle
    soup = BeautifulSoup(data, "html.parser")
    # Get the main div.
    main_div = soup.find('div', {'id': 'main'})

    # pre define the
    dict_list = []
    width = 0

    # loop over all the links
    for a in main_div.findAll('a'):
        print a
#        dict = {'href': u"" + a['href'], 'desc': u"" + a.find('img')['alt']}
        dict = {'href': u"" + a['href']}
        width = max(len(dict['href']), width)
        dict_list.append(dict)

    # Show the results just for nice print
    print ""
    print "Found %d items." % len(dict_list)
    for l in dict_list:
        print '{message: <{width}}'.format(message=l['desc'].encode('utf8'), width=width) + "\t" + l['href'].encode('utf8')

    print ""
    # Process...
    for l in dict_list:
        if not "produkt" in l['href'].encode('utf8'):
            continue
        processProduct(base_url, l)


def tools(request):

    latest_entry_list = Entry.objects.order_by('-creation_date')[:3]
    catagory_list = Category.objects.raw(
        "select cate.id, cate.title, cate.description, count(cate.id) as entrynum from blog_category as cate, blog_entry_categorys as enca where cate.id = enca.category_id group by cate.id")
    tag_list = Tag.objects.all()
    context = {'catagory_list': catagory_list,
               'latest_post': latest_entry_list, 'tag_list': tag_list}

    return render(request, 'tools.html', context)


def readdoc(request):

    print('124125')

    wordfile = request.FILES['wordfile']

    with open('/Users/valentine/Downloads/readbook1.docx', 'wb+') as destination:
        for chunk in wordfile.chunks():
            destination.write(chunk)

    document = Document(u'/Users/valentine/Downloads/readbook1.docx')
    # listp = [paragraph.text.encode('gb2312') for paragraph in document.paragraphs]

    counters = {}
    listp = document.paragraphs
    for i in listp:
        # print i.text
        items = i.text.split(' ')
        for item in items:
            if item in counters:
                counters[item] += 1
            else:
                counters[item] = 1

    # print "count of different word:"
    # print counters

    data = xlwt.Workbook()
    table = data.add_sheet('name', cell_overwrite_ok=True)
    line = 0
    for (k, v) in counters.items():
        print k, v
        table.write(line, 0, k)
        table.write(line, 1, v)
        line += 1

    data.save(u'/Users/valentine/Downloads/test3.xls')

    # print "Most popular word:"
    # print sorted([(counter,word) for word,counter in counters.items()],reverse=True)[0][1]

    def file_iterator(file_name, chunk_size=512):
        with open(file_name) as f:
            while True:
                c = f.read(chunk_size)
                if c:
                    yield c
                else:
                    break

    the_file_name = "test3.xls"
    response = StreamingHttpResponse(file_iterator(u'/Users/valentine/Downloads/test3.xls'))
    response['Content-Type'] = 'application/octet-stream'
    response['Content-Disposition'] = 'attachment;filename="{0}"'.format(the_file_name)

    return response
